import sys
import base64
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_dynamic_field(run, field_code):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def get_mermaid_image(mermaid_text, output_path):
    # Encode the mermaid text to base64
    encoded_text = base64.urlsafe_b64encode(mermaid_text.encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{encoded_text}"
    response = requests.get(url)
    if response.status_code == 200:
        with open(output_path, 'wb') as f:
            f.write(response.content)
        print(f"Downloaded {output_path} successfully.")
        return True
    print(f"Failed to download {output_path}. Status: {response.status_code}, URL length: {len(url)}")
    return False

# Create document
doc = Document()

# Thêm tiêu đề chính
title = doc.add_heading('TÀI LIỆU KỸ THUẬT: OLLAMA SMART ROUTER V4', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Thêm các Danh mục theo quy chuẩn
doc.add_heading('MỤC LỤC', level=1)
p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
add_dynamic_field(run_toc, 'TOC \\o "1-3" \\h \\z \\u')

doc.add_heading('DANH MỤC HÌNH ẢNH', level=1)
doc.add_paragraph('Hình 1: Sơ đồ khối luồng xử lý và định tuyến từ Prompt của người dùng')
doc.add_paragraph('Hình 2: Kiến trúc phân loại và kiểm tra chéo (Cross-Check) đa tác tử')

doc.add_heading('DANH MỤC BẢNG BIỂU', level=1)
doc.add_paragraph('Bảng 1: Danh sách Models (Inventory)')
doc.add_paragraph('Bảng 2: Tools Tích Hợp (MCP)')

doc.add_heading('DANH MỤC TỪ VIẾT TẮT', level=1)
doc.add_paragraph('- AI: Artificial Intelligence')
doc.add_paragraph('- MCP: Model Context Protocol')
doc.add_paragraph('- VRAM: Video Random Access Memory')
doc.add_paragraph('- TOC: Table of Contents')

doc.add_page_break()

doc.add_heading('1. Tổng quan các thảo luận', level=1)
doc.add_paragraph(
    'Dự án phát triển hệ thống Smart Router cho phép tự động định tuyến (route) câu hỏi (prompt) '
    'của người dùng đến đúng mô hình AI phù hợp đang chạy nội bộ (Ollama) nhằm tối ưu hiệu năng và tài nguyên. '
    'Người dùng viết prompt bình thường mà không cần tự khai báo model.'
)
doc.add_paragraph('Quá trình thảo luận đã chốt các quyết định kiến trúc phân bổ model sau:')
doc.add_paragraph('- Các yêu cầu Code (viết/debug/refactor) -> Qwen 2.5 Coder 32B', style='List Bullet')
doc.add_paragraph('- Các yêu cầu Suy luận (toán/logic/địa kỹ thuật) -> Nemotron 3.5 Lightning', style='List Bullet')
doc.add_paragraph('- Các yêu cầu Chung/Nhanh (phân tích, dịch, hỏi đáp) -> Gemma4 26B', style='List Bullet')

doc.add_heading('2. Kiến trúc hệ thống', level=1)
mermaid_1 = """%%{init: {'themeVariables': {'fontSize': '28px'}}}%%
flowchart TD
    U["Người dùng gửi prompt"] --> D{"Có @model?"}
    D -->|"@gemma"| G26["ask_gemma: Gemma4 26B (17GB)"]
    D -->|"@qwen"| Q["ask_qwen: Qwen Coder 32B (20GB)"]
    D -->|"@nemotron"| N["ask_nemotron: Nemotron 3.5 (25GB)"]
    D -->|"Không chỉ định"| A["ask_auto: Smart Router"]
    A --> C{"Phân loại prompt"}
    C -->|"code"| Q
    C -->|"reasoning"| N
    C -->|"general / quick"| G26
    style Q fill:#2563eb,color:#fff
    style N fill:#7c3aed,color:#fff
    style G26 fill:#059669,color:#fff
    style A fill:#f59e0b,color:#000
"""
doc.add_paragraph('Dưới đây là sơ đồ khối luồng xử lý và định tuyến từ Prompt của người dùng:')
if get_mermaid_image(mermaid_1, 'arch.png'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture('arch.png', width=Inches(6.0))
    p_caption = doc.add_paragraph('Hình 1: Sơ đồ khối luồng xử lý và định tuyến từ Prompt của người dùng')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.style = 'Caption'

doc.add_heading('3. Cơ chế Multi-Agent Cross-Check', level=1)
mermaid_2 = """%%{init: {'themeVariables': {'fontSize': '56px'}}}%%
flowchart TD
    P["Prompt"] --> K1{"Phân loại\\nbằng Keywords"}
    K1 -->|Code| M1["Primary: Qwen"]
    K1 -->|Reasoning| M2["Primary: Nemotron"]
    K1 -->|General/Quick| M3["Primary: Gemma 26B"]
    M1 -->|Draft| R["Reviewer: Gemma 26B"]
    M2 -->|Draft| R
    M3 -->|Draft| FINAL["Kết quả cuối cùng"]
    R -->|Đã duyệt| FINAL
    style M1 fill:#2563eb,color:#fff
    style M2 fill:#7c3aed,color:#fff
    style M3 fill:#059669,color:#fff
    style R fill:#10b981,color:#fff
"""
doc.add_paragraph('Kiến trúc phân loại và kiểm tra chéo (Cross-Check) đa tác tử để đảm bảo chất lượng trả lời cao nhất:')
if get_mermaid_image(mermaid_2, 'logic.png'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture('logic.png', height=Inches(6.5))
    p_caption2 = doc.add_paragraph('Hình 2: Kiến trúc phân loại và kiểm tra chéo (Cross-Check) đa tác tử')
    p_caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption2.style = 'Caption'

doc.add_heading('4. Danh sách Models (Inventory)', level=1)
p_tab1 = doc.add_paragraph('Bảng 1: Danh sách Models (Inventory)')
p_tab1.alignment = WD_ALIGN_PARAGRAPH.CENTER
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, name in enumerate(['Model', 'Params', 'VRAM', 'Context', 'Task Type', 'Tốc độ']):
    hdr_cells[i].text = name
    hdr_cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('Gemma4 26B', '25.8B', '17 GB', '262K', 'general/quick', '★★☆'),
    ('Qwen 2.5 Coder 32B', '32B', '20 GB', '128K', 'code', '★★☆'),
    ('Nemotron 3.5 Lightning', '42B', '25 GB', '128K', 'reasoning', '★☆☆')
]
for item in data:
    row_cells = table.add_row().cells
    for i, text in enumerate(item):
        row_cells[i].text = text

doc.add_heading('5. Tools Tích Hợp (MCP)', level=1)
p_tab2 = doc.add_paragraph('Bảng 2: Tools Tích Hợp (MCP)')
p_tab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr = table2.rows[0].cells
for i, name in enumerate(['Tool', 'Model', 'Trigger']):
    hdr[i].text = name
    hdr[i].paragraphs[0].runs[0].bold = True

tools = [
    ('ask_auto', 'Tự chọn', 'Mặc định — agent phân loại và kiểm duyệt (Cross-Check)'),
    ('ask_gemma', 'Gemma4 26B', '@gemma hoặc phân tích dài/ngắn'),
    ('ask_qwen', 'Qwen Coder 32B', '@qwen hoặc code tasks'),
    ('ask_nemotron', 'Nemotron 3.5', '@nemotron hoặc reasoning')
]
for item in tools:
    row = table2.add_row().cells
    for i, text in enumerate(item):
        row[i].text = text

doc.add_heading('6. Quản lý hệ thống tự động cài đặt', level=1)
doc.add_paragraph('Để dễ dàng đóng gói và chia sẻ, toàn bộ hệ thống đã được viết thành công cụ auto-installer (install.bat) thông minh gồm 9 bước tự động:')
steps = [
    'Cài đặt IDE (Antigravity IDE qua winget)',
    'Cài đặt môi trường Python 3.12',
    'Cài đặt Ollama engine',
    'Thiết lập MCP Smart Router',
    'Đăng ký Skills hỗ trợ Agent',
    'Cấu hình MCP Server JSON',
    'Tự động kéo mô hình Ollama (theo nhóm/tuỳ chọn)',
    'Cài đặt Thư viện Python theo 10 phân nhóm (BIM, Geotech, Data Science, Document, v.v)',
    'Cài đặt 35 Plugins IDE tự động theo 6 phân nhóm'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}", style='List Number')

# Thêm cấu hình updateFields tự động cập nhật TOC
element_updatefields = OxmlElement('w:updateFields')
element_updatefields.set(qn('w:val'), 'true')
doc.settings.element.append(element_updatefields)

# Header
section = doc.sections[0]
header = section.header
p_head = header.paragraphs[0]
p_head.text = "Phạm Ngọc Bảy - Liên hệ Zalo: 0972290107"
p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Footer
footer = section.footer
p_foot = footer.paragraphs[0]
p_foot.text = "Trang "
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_foot = p_foot.add_run()
add_dynamic_field(run_foot, 'PAGE')

# Save document
doc.save(r'C:\DEMO AI AGENT\ollama-smart-router\Ollama_Smart_Router_Documentation.docx')
doc.save(r'C:\DEMO AI AGENT\Ollama_Smart_Router_Documentation.docx')
print("Document generated successfully.")
